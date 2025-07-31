import time
import base64
import os
import threading
import keyboard
import speech_recognition as sr
from queue import Queue, Empty  
import pyttsx3
from google import genai
from google.genai import types
from docx import Document
from datetime import datetime


def save_report_to_word(report_text):
    doc = Document()
    doc.add_heading("Consultation Report", 0)
    
    for line in report_text.split('\n'):
        if line.strip():  # Skip empty lines
            doc.add_paragraph(line.strip())

    # Get the script's directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Build the filename with full path
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"Consultation_Report_{timestamp}.docx"
    full_path = os.path.join(script_dir, filename)
    
    doc.save(full_path)
    print(f"\n?? Report saved as: {full_path}")

def speech_to_text_continuous():
    audio_queue = Queue()
    stop_event = threading.Event()
    recognizer = sr.Recognizer()
    combined_text = []  # Store text from each chunk

    def record_audio():
        with sr.Microphone() as source:
            recognizer.adjust_for_ambient_noise(source)
            print("Recording... Press ESC to stop")
            
            while not stop_event.is_set():
                try:
                    audio = recognizer.listen(source, timeout=0.1, phrase_time_limit=10)
                    audio_queue.put(audio)
                except sr.WaitTimeoutError:
                    # Handle timeout between audio chunks
                    continue
                except Exception as e:
                    print(f"Recording error: {e}")
                    break

    def process_audio():
        while not stop_event.is_set() or not audio_queue.empty():
            try:
                audio = audio_queue.get(timeout=1)
                text = recognizer.recognize_google(audio)
                combined_text.append(text)
                print(f"Partial: {text}")
            except Empty:  # FIX: Use directly imported Empty
                continue
            except sr.UnknownValueError:
                pass  # Skip unrecognized audio
            except sr.RequestError as e:
                print("Skipping a segment due to network error.")
                continue
    # Start recording and processing threads
    record_thread = threading.Thread(target=record_audio, daemon=True)
    process_thread = threading.Thread(target=process_audio, daemon=True)
    
    record_thread.start()
    process_thread.start()
    
    input("Press ENTER to stop recording...\n")
    #keyboard.wait('esc')
    stop_event.set()
    print("\nStopping...")
    
    # Wait for threads to finish
    record_thread.join(timeout=2.0)
    process_thread.join(timeout=2.0)
    
    return " ".join(combined_text)


def read_text(text):
    """Reads the given text using a text-to-speech engine."""
    engine = pyttsx3.init()

    # Get available voices
    voices = engine.getProperty('voices')

    # Try to find a female voice, else default to the first available voice
    female_voice = None
    for voice in voices:
        if "female" in voice.name.lower() or "Zira" in voice.name or "Samantha" in voice.name:
            female_voice = voice.id
            break

    if female_voice:
        engine.setProperty('voice', female_voice)
    else:
        engine.setProperty('voice', voices[0].id)  # Default to first voice if no female voice is found

    engine.setProperty('rate', 150)  # Set speech rate
    engine.say(text)
    engine.runAndWait()


script_dir = os.path.dirname(os.path.abspath(__file__))
prompt_path = os.path.join(script_dir, "medical_prompt.txt")

with open(prompt_path, "r", encoding="utf-8") as file:
    medical_prompt = file.read()


def generate(user_input, chat_history):
    client = genai.Client(
        api_key=("AIzaSyC8ZXhnkI7iJJpYLMKm2zaEtMuuQqOEkXQ"), 
    )

    model = "gemini-2.0-flash"

    contents = []
    contents.append(types.Content(role="user",parts=[types.Part.from_text(text= medical_prompt)]))

    for message in chat_history:
        contents.append(types.Content(role=message["role"], parts=[types.Part.from_text(text=message["text"])]))

    contents.append(types.Content(role="user", parts=[types.Part.from_text(text=user_input)]))

    tools = [
        types.Tool(google_search=types.GoogleSearch())
    ]
    generate_content_config = types.GenerateContentConfig(
        temperature=1,
        top_p=0.95,
        top_k=40,
        max_output_tokens=8192,
        tools=tools,
        response_mime_type="text/plain",
    )

    response = ""
    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        print(chunk.text, end="")
        response += chunk.text

    return response

chat_history = []

result = speech_to_text_continuous()

response = generate(result, chat_history)
cleaned_response = response.replace("*", "")

save_report_to_word(cleaned_response)

chat_history.append({"role": "user", "text": result})
chat_history.append({"role": "model", "text": response})